"""DOCX template filler — detects yellow-highlighted fields and fills them."""

from __future__ import annotations

import io
import re
from datetime import date
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

from models import MergedData


def _na(value):
    """Return 'N/A' if the value is empty, None, or the string 'null'."""
    if not value or str(value).strip().lower() == "null":
        return "N/A"
    return value


def _format_certification_periods(data: MergedData) -> str:
    if not data.certification_periods:
        return "N/A"
    parts = []
    for period in data.certification_periods:
        parts.append(f"{period.start_date} - {period.end_date}")
    return "; ".join(parts)


def _ordinal(day: int) -> str:
    """Return day with ordinal suffix (1st, 2nd, 3rd, 4th, ...)."""
    if 11 <= day <= 13:
        return f"{day}th"
    suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suffix}"


def _today_formatted() -> str:
    """Return today's date as 'Month Dayth, YYYY' (e.g. 'March 3rd, 2026')."""
    today = date.today()
    return f"{today.strftime('%B')} {_ordinal(today.day)}, {today.year}"


# Maps normalized label text (before the colon) to a callable: (MergedData) -> str
FIELD_LABEL_MAP = {
    "patient name": lambda d: _na(d.patient_name),
    "date of birth": lambda d: _na(d.dob),
    "secondary/related diagnoses": lambda d: _na(d.secondary_diagnoses),
    "hospice noe date": lambda d: _na(d.noe_date),
    "hospice terminal diagnosis": lambda d: _na(d.terminal_diagnosis),
    "certification date(s)": lambda d: _format_certification_periods(d),
}

# The primary diagnosis label can vary (e.g. "Primary HPP Diagnosis", "Primary UWI Diagnosis").
# We match any label starting with "primary" and ending with "diagnosis".
_PRIMARY_DX_RE = re.compile(r"^primary\b.*\bdiagnosis$", re.IGNORECASE)

# Date line pattern: "Month Day, YYYY" or "Month Dayth, YYYY"
_DATE_LINE_RE = re.compile(
    r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}(st|nd|rd|th)?,?\s*\d{4}$"
)


def _run_is_yellow(run) -> bool:
    """Check if a run has yellow highlighting."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return False
    hl = rpr.find(qn("w:highlight"))
    if hl is None:
        return False
    return hl.get(qn("w:val")) == "yellow"


def _paragraph_is_yellow(para) -> bool:
    """Check if all non-empty runs in a paragraph are yellow-highlighted."""
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    return all(_run_is_yellow(r) for r in runs_with_text)


def _paragraph_has_yellow_run(para) -> bool:
    """Check if any run in the paragraph is yellow-highlighted."""
    return any(_run_is_yellow(r) for r in para.runs if r.text.strip())


def _is_date_paragraph(para) -> bool:
    """Check if this paragraph contains a date line like 'February 19th, 2026'."""
    text = para.text.strip()
    return bool(_DATE_LINE_RE.match(text))


def _extract_field_label(text: str) -> str | None:
    """Extract the label before the colon, if present."""
    if ":" not in text:
        return None
    label = text.split(":")[0].strip()
    return label if label else None


def _resolve_field_value(label: str, data: MergedData) -> str | None:
    """Look up a field label and return the resolved value, or None if unknown."""
    normalized = label.strip().lower()

    # Check exact match first
    resolver = FIELD_LABEL_MAP.get(normalized)
    if resolver:
        return resolver(data)

    # Check primary diagnosis pattern (e.g. "Primary HPP Diagnosis")
    if _PRIMARY_DX_RE.match(normalized):
        return _na(data.primary_diagnosis_code)

    return None


def _set_paragraph_text(para, new_text: str) -> None:
    """Replace all run text in a paragraph, preserving first run's formatting."""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def fill_docx_template(
    template_bytes: bytes,
    data: MergedData,
    addressee: str = "",
    company_name: str = "",
) -> bytes:
    """Fill a DOCX template by detecting yellow-highlighted fields.

    Parameters
    ----------
    template_bytes : bytes
        Raw bytes of the uploaded DOCX template.
    data : MergedData
        Merged extraction data from the 3 source documents.
    addressee : str
        Multi-line addressee (facility name, street, city/state/zip).
    company_name : str
        Company name to replace the yellow-highlighted company name in the body.

    Returns
    -------
    bytes
        The filled DOCX document as bytes.
    """
    doc = Document(io.BytesIO(template_bytes))
    paragraphs = doc.paragraphs

    # Split addressee into lines
    addressee_lines = [line.strip() for line in addressee.strip().splitlines() if line.strip()]

    # --- Pass 1: Find and fill date line ---
    for para in paragraphs:
        if _is_date_paragraph(para):
            _set_paragraph_text(para, _today_formatted())
            break

    # --- Pass 2: Find and fill addressee block ---
    # The addressee is a consecutive block of fully-yellow paragraphs before the "Re:" line.
    addressee_start = None
    addressee_end = None
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if text.startswith("Re:"):
            # Walk backwards to find the yellow block
            j = i - 1
            # Skip empty paragraphs
            while j >= 0 and not paragraphs[j].text.strip():
                j -= 1
            # Find consecutive yellow paragraphs
            end = j
            while j >= 0 and _paragraph_is_yellow(paragraphs[j]):
                j -= 1
            start = j + 1
            if start <= end:
                addressee_start = start
                addressee_end = end
            break

    if addressee_start is not None:
        addr_paras = list(range(addressee_start, addressee_end + 1))
        if addressee_lines:
            # Fill existing addressee paragraphs with new lines
            for idx, para_idx in enumerate(addr_paras):
                if idx < len(addressee_lines):
                    _set_paragraph_text(paragraphs[para_idx], addressee_lines[idx])
                else:
                    _set_paragraph_text(paragraphs[para_idx], "")

            # If user provided more lines than template has, append to the last paragraph
            if len(addressee_lines) > len(addr_paras):
                extra = " ".join(addressee_lines[len(addr_paras):])
                last_para = paragraphs[addressee_end]
                last_para.runs[0].text = last_para.runs[0].text + " " + extra
        else:
            # No addressee provided — replace with placeholder
            _set_paragraph_text(paragraphs[addressee_start], "[Addressee]")
            for para_idx in addr_paras[1:]:
                _set_paragraph_text(paragraphs[para_idx], "")

    # --- Pass 3: Find and fill yellow field lines (contain ":") ---
    for para in paragraphs:
        if not _paragraph_is_yellow(para):
            continue
        text = para.text.strip()
        label = _extract_field_label(text)
        if label is None:
            continue
        value = _resolve_field_value(label, data)
        if value is not None:
            _set_paragraph_text(para, f"{label}: {value}")

    # --- Pass 4: Find and replace yellow company name run in body paragraphs ---
    if company_name:
        for para in paragraphs:
            if not _paragraph_has_yellow_run(para):
                continue
            # Skip paragraphs that are fully yellow (addressee or field lines)
            if _paragraph_is_yellow(para):
                continue
            # This is a mixed paragraph with some yellow runs — replace yellow runs
            for run in para.runs:
                if _run_is_yellow(run) and run.text.strip():
                    run.text = company_name

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def scan_template_placeholders(template_bytes: bytes) -> list[dict]:
    """Scan a DOCX template and return info about detected yellow fields.

    Returns a list of dicts with keys: 'field', 'type', 'status'.
    """
    doc = Document(io.BytesIO(template_bytes))
    results = []

    for para in doc.paragraphs:
        if _is_date_paragraph(para):
            results.append({
                "field": "Date",
                "type": "Auto-fill",
                "status": "Will be set to today's date",
            })
            continue

        if not _paragraph_has_yellow_run(para):
            continue

        text = para.text.strip()

        # Field line with colon
        label = _extract_field_label(text)
        if label and _paragraph_is_yellow(para):
            normalized = label.strip().lower()
            is_known = normalized in FIELD_LABEL_MAP or bool(_PRIMARY_DX_RE.match(normalized))
            results.append({
                "field": label,
                "type": "Data field",
                "status": "Mapped" if is_known else "Unknown",
            })
            continue

        # Fully yellow paragraph without colon (addressee)
        if _paragraph_is_yellow(para):
            results.append({
                "field": text[:50],
                "type": "Addressee",
                "status": "Will use addressee input",
            })
            continue

        # Mixed paragraph with some yellow (company name)
        yellow_text = " ".join(r.text for r in para.runs if _run_is_yellow(r) and r.text.strip())
        if yellow_text:
            results.append({
                "field": yellow_text[:50],
                "type": "Company name",
                "status": "Will use company name input",
            })

    return results
